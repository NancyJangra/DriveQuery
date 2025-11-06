"""
Document Processing Service
============================
Handles uploading, reading, and processing documents (PDF, DOCX, TXT).

For beginners: This is just using libraries to read files!
- PyPDF2 reads PDFs
- python-docx reads Word files
- We just extract text and save it
"""

import os
import uuid
from typing import List, Tuple, Optional
import logging
from pathlib import Path
import json

# Document processing libraries
import PyPDF2
from docx import Document

# Try to import optional dependencies
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_openai import OpenAIEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain.schema import Document as LangChainDocument
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    print("⚠️  LangChain not installed - document search will use simple text matching")

from app.core.config import settings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentService:
    """
    Service to handle document upload and processing
    Think of this as your "file handler"
    """
    
    def __init__(self):
        """Initialize the document service"""
        global LANGCHAIN_AVAILABLE
        self.upload_dir = "uploads"
        self.vector_store_dir = "vector_store"
        self.document_store_path = os.path.join(self.vector_store_dir, "documents.json")
        
        # Create directories if they don't exist
        os.makedirs(self.upload_dir, exist_ok=True)
        os.makedirs(self.vector_store_dir, exist_ok=True)
        
        # Initialize embeddings and vector store if LangChain is available
        if LANGCHAIN_AVAILABLE:
            try:
                self.embeddings = OpenAIEmbeddings(
                    openai_api_key=settings.OPENAI_API_KEY
                )
                self.vector_store: Optional[FAISS] = None
                self._load_vector_store()
                logger.info("✅ Document Service initialized with FAISS")
            except Exception as e:
                logger.warning(f"⚠️  Could not initialize FAISS: {e}")
                self.vector_store = None
                LANGCHAIN_AVAILABLE = False
        else:
            self.vector_store = None
            logger.info("✅ Document Service initialized (without FAISS - using simple text search)")
        
        # Load document metadata
        self.documents_metadata = self._load_documents_metadata()
    
    def _load_documents_metadata(self) -> dict:
        """Load document metadata from JSON file"""
        if os.path.exists(self.document_store_path):
            try:
                with open(self.document_store_path, 'r') as f:
                    return json.load(f)
            except:
                return {}
        return {}
    
    def _save_documents_metadata(self):
        """Save document metadata to JSON file"""
        with open(self.document_store_path, 'w') as f:
            json.dump(self.documents_metadata, f, indent=2)
    
    def _load_vector_store(self):
        """Load existing vector store if available"""
        vector_store_path = os.path.join(self.vector_store_dir, "faiss_index")
        if os.path.exists(vector_store_path):
            try:
                self.vector_store = FAISS.load_local(
                    vector_store_path, 
                    self.embeddings,
                    allow_dangerous_deserialization=True
                )
                logger.info("Loaded existing vector store")
            except Exception as e:
                logger.warning(f"Could not load vector store: {e}")
                self.vector_store = None
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF file
        
        Returns: (text_content, number_of_pages)
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                logger.info(f"Extracted {len(text)} characters from {num_pages} pages")
                return text, num_pages
                
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise Exception(f"Failed to read PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from Word document (.docx)
        
        Returns: text_content
        """
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            raise Exception(f"Failed to read DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from plain text file
        
        Returns: text_content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            logger.info(f"Extracted {len(text)} characters from TXT")
            return text
            
        except Exception as e:
            logger.error(f"Error reading TXT: {str(e)}")
            raise Exception(f"Failed to read TXT: {str(e)}")
    
    def process_document(self, file_path: str, filename: str) -> dict:
        """
        Process uploaded document and add to vector store
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
        
        Returns:
            Dictionary with document info
        """
        
        # Get file extension
        file_ext = Path(filename).suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            text, pages = self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            text = self.extract_text_from_docx(file_path)
            pages = None
        elif file_ext == '.txt':
            text = self.extract_text_from_txt(file_path)
            pages = None
        else:
            raise Exception(f"Unsupported file type: {file_ext}")
        
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        
        # Store document metadata
        self.documents_metadata[doc_id] = {
            "filename": filename,
            "text": text,
            "pages": pages,
            "chunks": 0
        }
        
        # If LangChain is available, use advanced processing
        if LANGCHAIN_AVAILABLE:
            try:
                # Split text into chunks
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len,
                )
                
                chunks = text_splitter.split_text(text)
                logger.info(f"Split document into {len(chunks)} chunks")
                
                # Create LangChain documents with metadata
                documents = [
                    LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": filename,
                            "doc_id": doc_id,
                            "chunk_index": i
                        }
                    )
                    for i, chunk in enumerate(chunks)
                ]
                
                # Add to vector store (FAISS)
                if self.vector_store is None:
                    self.vector_store = FAISS.from_documents(documents, self.embeddings)
                else:
                    self.vector_store.add_documents(documents)
                
                # Save vector store to disk
                vector_store_path = os.path.join(self.vector_store_dir, "faiss_index")
                self.vector_store.save_local(vector_store_path)
                logger.info("✅ Vector store updated and saved")
                
                self.documents_metadata[doc_id]["chunks"] = len(chunks)
            except Exception as e:
                logger.warning(f"⚠️  Could not create embeddings: {e}. Document saved but search will be limited.")
                # Document is still saved, just without advanced search
        
        # Save metadata
        self._save_documents_metadata()
        
        return {
            "doc_id": doc_id,
            "filename": filename,
            "pages": pages,
            "chunks": self.documents_metadata[doc_id]["chunks"],
            "total_chars": len(text)
        }
    
    def search_documents(self, query: str, top_k: int = 3) -> List[str]:
        """
        Search uploaded documents for relevant information
        
        This uses FAISS if available, otherwise uses simple text search
        
        Args:
            query: User's question
            top_k: Number of relevant chunks to return
        
        Returns:
            List of relevant text chunks
        """
        
        # If FAISS is available and initialized, use it
        if LANGCHAIN_AVAILABLE and self.vector_store is not None:
            try:
                results = self.vector_store.similarity_search(query, k=top_k)
                relevant_texts = [doc.page_content for doc in results]
                logger.info(f"✅ Found {len(relevant_texts)} relevant chunks using FAISS")
                return relevant_texts
            except Exception as e:
                logger.error(f"❌ Error with FAISS search: {str(e)}")
                # Fall back to simple search
        
        # Simple text search (fallback)
        logger.info("Using simple text search (FAISS not available)")
        relevant_texts = []
        
        for doc_id, metadata in self.documents_metadata.items():
            if 'text' in metadata:
                text = metadata['text']
                # Simple keyword matching
                query_words = query.lower().split()
                text_lower = text.lower()
                
                # Check if query words appear in text
                matches = sum(1 for word in query_words if word in text_lower)
                if matches > 0:
                    # Extract relevant portion (approximate)
                    relevant_texts.append(text[:1000])  # First 1000 chars
                    
        return relevant_texts[:top_k]
    
    def get_all_documents(self) -> List[str]:
        """
        Get list of all uploaded documents
        
        Returns:
            List of filenames
        """
        try:
            files = os.listdir(self.upload_dir)
            return [f for f in files if not f.startswith('.')]
        except Exception as e:
            logger.error(f"Error listing documents: {str(e)}")
            return []

# Create singleton instance
document_service = DocumentService()
